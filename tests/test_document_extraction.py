from io import BytesIO

import pytest
from docx import Document
from pypdf import PdfWriter
from reportlab.pdfgen import canvas

from candidature_emploi.infrastructure.document_extraction import (
    DocumentBlock,
    DocumentExtractionError,
    extract_document,
)
from candidature_emploi.infrastructure import document_extraction


def make_text_pdf() -> bytes:
    output = BytesIO()
    pdf = canvas.Canvas(output)
    pdf.drawString(72, 780, "Camille Exemple")
    pdf.drawString(72, 760, "Profil data engineer avec Python, SQL et Airflow.")
    pdf.drawString(72, 740, "Experience professionnelle et formation detaillees.")
    pdf.save()
    return output.getvalue()


def make_docx() -> bytes:
    output = BytesIO()
    document = Document()
    document.add_heading("Camille Exemple", level=1)
    document.add_paragraph("Profil data engineer avec Python, SQL et Airflow.")
    document.add_heading("Compétences", level=2)
    document.add_paragraph("Python, SQL, Airflow")
    document.save(output)
    return output.getvalue()


def make_blank_pdf() -> bytes:
    output = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.write(output)
    return output.getvalue()


def test_extracts_textual_pdf() -> None:
    result = extract_document("cv.pdf", make_text_pdf())

    assert result.source_format == "pdf"
    assert result.page_count == 1
    assert "Camille Exemple" in result.text


def test_extracts_docx() -> None:
    result = extract_document("cv.docx", make_docx())

    assert result.source_format == "docx"
    assert "Python, SQL, Airflow" in result.text


def test_rejects_corrupted_docx_with_a_user_facing_error() -> None:
    with pytest.raises(DocumentExtractionError, match="illisible ou endommagé"):
        extract_document("cv.docx", b"not-a-docx-file" * 10)


def test_rejects_scanned_or_blank_pdf() -> None:
    with pytest.raises(DocumentExtractionError, match="OCR"):
        extract_document("scan.pdf", make_blank_pdf())


def test_uses_local_ocr_when_pdf_text_is_unreliable(monkeypatch: pytest.MonkeyPatch) -> None:
    content = make_text_pdf()
    monkeypatch.setattr(document_extraction, "_pdf_text_needs_ocr", lambda text: True)
    monkeypatch.setattr(
        document_extraction,
        "_extract_pdf_with_ocr_blocks",
        lambda content, page_count: (
            "Camille Exemple\nCompétences\nPython, SQL, Airflow, PostgreSQL",
            (
                DocumentBlock(
                    text="Compétences",
                    page=1,
                    column="main",
                    x0=10,
                    y0=10,
                    x1=100,
                    y1=20,
                    source="ocr",
                ),
            ),
        ),
    )

    document = extract_document("cv.pdf", content)

    assert document.ocr_applied is True
    assert "Compétences" in document.text
    assert document.blocks[0].source == "ocr"


def test_orders_two_column_layout_left_then_right() -> None:
    blocks = [
        (10.0, 10.0, 190.0, 20.0, "Compétences"),
        (210.0, 10.0, 390.0, 20.0, "Expériences"),
        (10.0, 30.0, 190.0, 40.0, "Python"),
        (210.0, 30.0, 390.0, 40.0, "Data engineer"),
    ]

    ordered = document_extraction._order_layout_blocks(blocks, page_width=400.0)

    assert [block[4] for block in ordered] == ["Compétences", "Python", "Expériences", "Data engineer"]


def test_assigns_pdf_lines_to_independent_columns() -> None:
    blocks = [
        DocumentBlock("Compétences", 1, "main", 10, 10, 180, 20, "pdf"),
        DocumentBlock("Python", 1, "main", 10, 30, 180, 40, "pdf"),
        DocumentBlock("Expériences", 1, "main", 220, 10, 390, 20, "pdf"),
        DocumentBlock("Data engineer", 1, "main", 220, 30, 390, 40, "pdf"),
    ]

    assigned = document_extraction._assign_columns(blocks, page_width=400)

    assert [block.column for block in assigned] == [
        "left",
        "left",
        "right",
        "right",
    ]


def test_assigns_short_right_heading_from_its_left_alignment() -> None:
    blocks = [
        DocumentBlock("Langues", 1, "main", 50, 10, 180, 20, "ocr"),
        DocumentBlock("Français", 1, "main", 50, 30, 220, 40, "ocr"),
        DocumentBlock("Expériences", 1, "main", 350, 10, 470, 20, "ocr"),
        DocumentBlock("Description longue", 1, "main", 350, 30, 900, 40, "ocr"),
    ]

    assigned = document_extraction._assign_columns(blocks, page_width=1000)

    assert [block.column for block in assigned] == [
        "left",
        "left",
        "right",
        "right",
    ]


def test_requests_ocr_when_core_sections_are_not_readable() -> None:
    assert document_extraction._pdf_text_needs_ocr("Camille Exemple\nLangues\nFrançais") is True
    assert document_extraction._pdf_text_needs_ocr(
        "Compétences\nPython, SQL\nExpériences\nData engineer\nFormations\nMaster Data"
    ) is False


def test_rejects_zero_confidence_ocr_artifacts() -> None:
    assert document_extraction._ocr_word_is_reliable("0") is False
    assert document_extraction._ocr_word_is_reliable("71") is True


def test_rejects_unsupported_format() -> None:
    with pytest.raises(DocumentExtractionError, match="Format non pris en charge"):
        extract_document("cv.txt", b"Texte suffisamment long pour un faux CV.")


def test_rejects_oversized_file() -> None:
    with pytest.raises(DocumentExtractionError, match="5 Mo"):
        extract_document("cv.pdf", b"x" * (5 * 1024 * 1024 + 1))
