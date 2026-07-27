"""Extraction locale de texte depuis un PDF textuel ou un DOCX."""

from __future__ import annotations

from dataclasses import dataclass, replace
from io import BytesIO
from math import ceil
from pathlib import Path
import unicodedata
from zipfile import BadZipFile

from docx import Document
from pypdf import PdfReader
from pypdf.errors import PdfReadError

MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024
MIN_MEANINGFUL_TEXT_LENGTH = 40
SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


class DocumentExtractionError(ValueError):
    """Erreur montrable à l'utilisateur sans divulguer de détail technique."""


@dataclass(frozen=True, slots=True)
class DocumentBlock:
    """Ligne de CV localisée, conservée seulement pendant la session."""

    text: str
    page: int
    column: str
    x0: float
    y0: float
    x1: float
    y1: float
    source: str


@dataclass(frozen=True, slots=True)
class ExtractedDocument:
    """Résultat temporaire d'une extraction locale."""

    text: str
    source_format: str
    page_count: int | None = None
    ocr_applied: bool = False
    has_embedded_image: bool = False
    blocks: tuple[DocumentBlock, ...] = ()


def extract_document(filename: str, content: bytes) -> ExtractedDocument:
    """Valide puis extrait un document sans l'enregistrer sur disque."""

    extension = Path(filename).suffix.casefold()
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentExtractionError(
            "Format non pris en charge. Chargez un fichier PDF textuel ou DOCX."
        )
    if not content:
        raise DocumentExtractionError("Le fichier chargé est vide.")
    if len(content) > MAX_FILE_SIZE_BYTES:
        raise DocumentExtractionError("Le fichier dépasse la limite de 5 Mo.")

    if extension == ".pdf":
        return _extract_pdf(content)
    return _extract_docx(content)


def _extract_pdf(content: bytes) -> ExtractedDocument:
    try:
        reader = PdfReader(BytesIO(content))
        page_text = [(page.extract_text() or "").strip() for page in reader.pages]
    except (PdfReadError, OSError, ValueError) as exc:
        raise DocumentExtractionError(
            "Le PDF est illisible ou endommagé."
        ) from exc

    text = _clean_text("\n\n".join(filter(None, page_text)))
    page_count = len(reader.pages)
    has_embedded_image = any(_page_has_image(page) for page in reader.pages)
    layout_text, layout_blocks = _extract_pdf_layout(content)
    if len(layout_text) >= MIN_MEANINGFUL_TEXT_LENGTH:
        text = layout_text
    if _pdf_text_needs_ocr(text):
        ocr_text, ocr_blocks = _extract_pdf_with_ocr_blocks(content, page_count)
        if len(ocr_text) >= MIN_MEANINGFUL_TEXT_LENGTH:
            return ExtractedDocument(
                text=ocr_text,
                source_format="pdf",
                page_count=page_count,
                ocr_applied=True,
                has_embedded_image=has_embedded_image,
                blocks=ocr_blocks,
            )
    if len(text) < MIN_MEANINGFUL_TEXT_LENGTH:
        raise DocumentExtractionError(
            "Aucun texte exploitable n’a été détecté, y compris après l’OCR local."
        )
    return ExtractedDocument(
        text=text,
        source_format="pdf",
        page_count=page_count,
        has_embedded_image=has_embedded_image,
        blocks=layout_blocks,
    )


def _extract_docx(content: bytes) -> ExtractedDocument:
    try:
        document = Document(BytesIO(content))
    except (BadZipFile, OSError, ValueError) as exc:
        raise DocumentExtractionError(
            "Le document DOCX est illisible ou endommagé."
        ) from exc

    blocks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text for cell in row.cells))
    text = _clean_text("\n".join(blocks))
    if len(text) < MIN_MEANINGFUL_TEXT_LENGTH:
        raise DocumentExtractionError(
            "Le document ne contient pas assez de texte exploitable."
        )
    return ExtractedDocument(text=text, source_format="docx")


def _pdf_text_needs_ocr(text: str) -> bool:
    """Évite l'OCR coûteux lorsque le PDF possède déjà un texte exploitable."""

    if len(text) < MIN_MEANINGFUL_TEXT_LENGTH:
        return True
    unreadable_markers = text.count("�") + text.count("□") + text.count("■")
    if unreadable_markers / max(len(text), 1) > 0.01:
        return True
    normalized = _ascii_text(text)
    core_sections = ("experience", "competence", "formation")
    return sum(section in normalized for section in core_sections) < 2


def _extract_pdf_with_ocr(content: bytes, page_count: int) -> str:
    """Rend les pages en mémoire et les lit localement avec Tesseract."""

    text, _ = _extract_pdf_with_ocr_blocks(content, page_count)
    return text


def _extract_pdf_with_ocr_blocks(
    content: bytes,
    page_count: int,
) -> tuple[str, tuple[DocumentBlock, ...]]:
    """Conserve les lignes et coordonnées renvoyées par l'OCR local."""

    if page_count > 5:
        return "", ()
    try:
        import pymupdf
        import pytesseract
        from PIL import Image
    except ImportError:
        return "", ()
    try:
        document = pymupdf.open(stream=content, filetype="pdf")
        page_blocks: list[DocumentBlock] = []
        for page_index, page in enumerate(document, start=1):
            pixmap = page.get_pixmap(dpi=200, alpha=False)
            image = Image.open(BytesIO(pixmap.tobytes("png")))
            data = pytesseract.image_to_data(
                image,
                lang="fra+eng",
                config="--psm 3",
                output_type=pytesseract.Output.DICT,
            )
            grouped: dict[tuple[int, int, int], list[int]] = {}
            for index, word in enumerate(data["text"]):
                if str(word).strip() and _ocr_word_is_reliable(data["conf"][index]):
                    key = (
                        int(data["block_num"][index]),
                        int(data["par_num"][index]),
                        int(data["line_num"][index]),
                    )
                    grouped.setdefault(key, []).append(index)
            raw_lines: list[DocumentBlock] = []
            for indexes in grouped.values():
                text = " ".join(str(data["text"][index]).strip() for index in indexes)
                left = min(int(data["left"][index]) for index in indexes)
                top = min(int(data["top"][index]) for index in indexes)
                right = max(
                    int(data["left"][index]) + int(data["width"][index])
                    for index in indexes
                )
                bottom = max(
                    int(data["top"][index]) + int(data["height"][index])
                    for index in indexes
                )
                raw_lines.append(
                    DocumentBlock(
                        text=text,
                        page=page_index,
                        column="main",
                        x0=float(left),
                        y0=float(top),
                        x1=float(right),
                        y1=float(bottom),
                        source="ocr",
                    )
                )
            page_blocks.extend(_assign_columns(raw_lines, float(image.width)))
        document.close()
    except Exception:
        return "", ()
    blocks = tuple(_order_document_blocks(page_blocks))
    return _clean_text("\n".join(block.text for block in blocks)), blocks


def _ocr_word_is_reliable(confidence: object) -> bool:
    """Écarte les formes OCR que Tesseract lui-même juge non fiables."""

    try:
        return float(confidence) >= 20
    except (TypeError, ValueError):
        return False


def _extract_pdf_layout_text(content: bytes) -> str:
    """Reconstruit l'ordre de lecture des PDF à deux colonnes, en mémoire."""

    text, _ = _extract_pdf_layout(content)
    return text


def _extract_pdf_layout(content: bytes) -> tuple[str, tuple[DocumentBlock, ...]]:
    """Extrait les lignes PDF en conservant leur géométrie."""

    try:
        import pymupdf
    except ImportError:
        return "", ()
    try:
        document = pymupdf.open(stream=content, filetype="pdf")
        extracted_blocks: list[DocumentBlock] = []
        for page_index, page in enumerate(document, start=1):
            raw_lines: list[DocumentBlock] = []
            page_dict = page.get_text("dict")
            for raw_block in page_dict.get("blocks", []):
                if raw_block.get("type") != 0:
                    continue
                for line in raw_block.get("lines", []):
                    text = "".join(
                        str(span.get("text", "")) for span in line.get("spans", [])
                    ).strip()
                    if not text:
                        continue
                    x0, y0, x1, y1 = (float(value) for value in line["bbox"])
                    raw_lines.append(
                        DocumentBlock(
                            text=text,
                            page=page_index,
                            column="main",
                            x0=x0,
                            y0=y0,
                            x1=x1,
                            y1=y1,
                            source="pdf",
                        )
                    )
            extracted_blocks.extend(_assign_columns(raw_lines, float(page.rect.width)))
        document.close()
    except Exception:
        return "", ()
    blocks = tuple(_order_document_blocks(extracted_blocks))
    return _clean_text("\n".join(block.text for block in blocks)), blocks


def _assign_columns(
    blocks: list[DocumentBlock],
    page_width: float,
) -> list[DocumentBlock]:
    """Identifie un flux simple ou deux colonnes sans perdre les coordonnées."""

    if not blocks:
        return []
    boundary = _column_boundary(blocks, page_width)
    if boundary is None:
        return [replace(block, column="main") for block in blocks]

    assigned: list[DocumentBlock] = []
    for block in blocks:
        width = block.x1 - block.x0
        if block.x0 <= page_width * 0.15 and width >= page_width * 0.75:
            column = "full"
        else:
            column = "left" if block.x0 < boundary else "right"
        assigned.append(replace(block, column=column))
    return assigned


def _column_boundary(
    blocks: list[DocumentBlock],
    page_width: float,
) -> float | None:
    """Trouve deux alignements de départ suffisamment peuplés et séparés."""

    positions = sorted(block.x0 for block in blocks)
    minimum_cluster = max(2, ceil(len(positions) * 0.18))
    candidates: list[tuple[float, int]] = []
    for index in range(minimum_cluster, len(positions) - minimum_cluster + 1):
        gap = positions[index] - positions[index - 1]
        candidates.append((gap, index))
    if not candidates:
        return None
    gap, index = max(candidates)
    if gap < page_width * 0.08:
        return None
    return (positions[index - 1] + positions[index]) / 2


def _order_document_blocks(blocks: list[DocumentBlock]) -> list[DocumentBlock]:
    """Ordonne chaque flux indépendamment pour éviter les croisements de colonnes."""

    column_order = {"full": 0, "main": 0, "left": 1, "right": 2}
    return sorted(
        blocks,
        key=lambda block: (
            block.page,
            column_order.get(block.column, 3),
            block.y0,
            block.x0,
        ),
    )


def _order_layout_blocks(
    blocks: list[tuple[float, float, float, float, str]],
    page_width: float,
) -> list[tuple[float, float, float, float, str]]:
    """Préserve les en-têtes puis lit la colonne gauche avant la droite."""

    midpoint = page_width / 2
    left = [block for block in blocks if block[2] <= midpoint]
    right = [block for block in blocks if block[0] >= midpoint]
    spanning = [block for block in blocks if block not in left and block not in right]
    sort_blocks = lambda values: sorted(values, key=lambda block: (block[1], block[0]))
    if len(left) < 2 or len(right) < 2:
        return sort_blocks(blocks)
    first_column_y = min(block[1] for block in left + right)
    headers = [block for block in spanning if block[3] <= first_column_y]
    footer = [block for block in spanning if block not in headers]
    return sort_blocks(headers) + sort_blocks(left) + sort_blocks(right) + sort_blocks(footer)


def _page_has_image(page: object) -> bool:
    """Détecte seulement la présence d'une image, sans l'extraire ni la conserver."""

    try:
        resources = page.get("/Resources", {})
        if hasattr(resources, "get_object"):
            resources = resources.get_object()
        xobjects = resources.get("/XObject", {}) if resources else {}
        if hasattr(xobjects, "get_object"):
            xobjects = xobjects.get_object()
        return any(
            (item.get_object() if hasattr(item, "get_object") else item).get("/Subtype") == "/Image"
            for item in xobjects.values()
        )
    except (AttributeError, TypeError):
        return False


def _ascii_text(value: str) -> str:
    value = unicodedata.normalize("NFKD", value)
    return "".join(character for character in value if not unicodedata.combining(character)).casefold()


def _clean_text(text: str) -> str:
    lines = [" ".join(line.split()) for line in text.replace("\x00", "").splitlines()]
    return "\n".join(line for line in lines if line).strip()
